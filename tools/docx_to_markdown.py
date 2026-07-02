#!/usr/bin/env python
"""
ARVES corpus renderer — .docx -> Markdown (v1.0.1 Foundation Improvement).

Renders the frozen specification corpus (the ~50 root .docx files) into a committed,
grep-/diff-/PR-able Markdown MIRROR under `spec-markdown/`. This is a FORMAT conversion
only: content is preserved, not changed.

IMPORTANT (ED-001): the .docx files remain the AUTHORITATIVE frozen corpus. The generated
Markdown is a faithful, regenerable *rendering* for readability, version control, line
review, search, and AI-agent consumption — it is NOT the source of record and MUST NOT be
hand-edited (regenerate instead). Byte-affecting spec change still goes through CCP/Amendment.

Order-preserving: iterates the document body so paragraphs and tables interleave correctly.
Maps Word heading styles to Markdown headings, list paragraphs to bullets, tables to GFM
tables, and bold/italic runs to Markdown emphasis.

Usage: python tools/docx_to_markdown.py           # convert all root *.docx -> spec-markdown/
       python tools/docx_to_markdown.py --check    # verify parity only (no write), exit 1 on drift
"""
import glob
import os
import re
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "spec-markdown")


def iter_block_items(doc):
    """Yield Paragraph and Table objects in true document order."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def run_text(par):
    """Concatenate runs, applying bold/italic as Markdown emphasis."""
    parts = []
    for run in par.runs:
        t = run.text
        if not t:
            continue
        if run.bold and run.italic:
            t = f"***{t}***"
        elif run.bold:
            t = f"**{t}**"
        elif run.italic:
            t = f"*{t}*"
        parts.append(t)
    return "".join(parts) if parts else par.text


def is_list(par):
    p = par._p
    return p.pPr is not None and p.pPr.numPr is not None


def para_md(par):
    style = (par.style.name if par.style else "") or ""
    s = style.lower()
    text = run_text(par).strip()
    if not text:
        return ""
    if s.startswith("heading"):
        digits = re.sub(r"\D", "", s)
        lvl = min(max(int(digits), 1), 6) if digits else 1
        return "#" * lvl + " " + text
    if s.startswith("title"):
        return "# " + text
    if s.startswith("subtitle"):
        return "## " + text
    if "list" in s or is_list(par):
        return "- " + text
    return text


def cell_text(cell):
    return " ".join(p.text.strip() for p in cell.paragraphs).replace("|", "\\|").strip()


def table_md(tbl):
    rows = list(tbl.rows)
    if not rows:
        return ""
    header = [cell_text(c) for c in rows[0].cells]
    lines = ["| " + " | ".join(header) + " |",
             "| " + " | ".join("---" for _ in header) + " |"]
    for r in rows[1:]:
        lines.append("| " + " | ".join(cell_text(c) for c in r.cells) + " |")
    return "\n".join(lines)


def convert(path):
    doc = Document(path)
    blocks = []
    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            md = para_md(item)
        else:
            md = table_md(item)
        if md:
            blocks.append(md)
    name = os.path.basename(path)
    banner = (f"> **Rendered from `{name}`** — a faithful Markdown mirror of the FROZEN "
              f"specification corpus (authoritative source of record: the `.docx`). "
              f"Format conversion only; content unchanged. **Do not hand-edit** — regenerate "
              f"via `python tools/docx_to_markdown.py`. Spec changes go through CCP/Amendment.\n")
    body = "\n\n".join(blocks)
    return banner + "\n" + body + "\n"


def docx_wordcount(path):
    doc = Document(path)
    words = 0
    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            words += len(item.text.split())
        else:
            for row in item.rows:
                for cell in row.cells:
                    words += len(cell_text(cell).split())
    return words


def main():
    check = "--check" in sys.argv
    files = sorted(glob.glob(os.path.join(ROOT, "*.docx")))
    if not files:
        print("no .docx at repo root", file=sys.stderr)
        return 1
    if not check:
        os.makedirs(OUT, exist_ok=True)
    ok = 0
    drift = 0
    for path in files:
        name = os.path.basename(path)
        md = convert(path)
        md_words = len(re.sub(r"[#*|>`-]", " ", md).split())
        src_words = docx_wordcount(path)
        # parity: rendered markdown should retain ~all source words (allow markup slack)
        parity = "OK" if md_words >= int(src_words * 0.95) else "LOW"
        if parity != "OK":
            drift += 1
        out_path = os.path.join(OUT, name[:-5] + ".md")
        if not check:
            with open(out_path, "w", encoding="utf-8", newline="\n") as f:
                f.write(md)
        ok += 1
        print(f"  {parity:3}  {name:60}  src_words={src_words:6}  md_words={md_words:6}")
    print(f"\n{ok}/{len(files)} rendered; {drift} parity-LOW.")
    if check and drift:
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())

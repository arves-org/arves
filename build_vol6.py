# -*- coding: utf-8 -*-
import docx
from docx.shared import Pt

FN = "ARVES_OS_Volume_6_Certification_Review_Manual_v1.docx"
d = docx.Document()

def H(t, lvl=1):
    d.add_heading(t, level=lvl)

def P(t=""):
    return d.add_paragraph(t)

def B(t):
    return d.add_paragraph(t, style='List Bullet')

def I(t):
    p = d.add_paragraph()
    r = p.add_run(t)
    r.italic = True
    return p

def table(headers, rows):
    t = d.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hc = t.rows[0].cells
    for i, h in enumerate(headers):
        hc[i].text = h
        for para in hc[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for row in rows:
        rc = t.add_row().cells
        for i, val in enumerate(row):
            rc[i].text = val
    return t

# ---------------- TITLE ----------------
d.add_paragraph('ARVES OS Volume 6 - Certification & Review Manual v1.0', style='Title')
P('STATUS: CERTIFICATION CONSTITUTION (INDEPENDENT VERDICT ON CORRECTNESS AND READINESS)')
P('SCOPE: Scenario Conformance Framework in depth; certification levels L1-L4 + Certified Product; '
  'the Independent Architecture Review process; conformance / performance / enterprise-readiness checklists; '
  'Independent Runtime A/B goals; certified ecosystem and product goals; versioning of conformance suites against spec versions.')
P('SPECIFICATION ERA: FROZEN as of 2026-07-01. IMPLEMENTATION ERA: IN PROGRESS. '
  'This volume certifies implementations against the frozen standard; it never modifies the standard.')

# ---------------- PART 1 ----------------
H('Part 1 - Purpose and Position in the Dependency Chain')
P('Volume 6 is the certification and review authority of the ARVES corpus. Where earlier volumes DEFINE '
  '(Theory, Specification, Contracts, Behaviour) and the Scenario Conformance Framework DECLARES what correctness '
  'means, this volume answers a single question: has a given implementation earned the right to call itself an '
  'ARVES runtime? It does so through two independent instruments - the Scenario Conformance Framework (mechanical, '
  'property-based verdicts) and the Independent Architecture Review (adversarial, human-judged verdicts).')
P('The dependency chain is never reversed:')
B('Theory -> Specification -> Contracts -> Behaviour -> Conformance -> Implementation.')
P('Certification sits at the Conformance -> Implementation boundary. An implementation PROVES the specification; '
  'it never changes it. A failed certification is a defect in the implementation (or, at most, a discovered gap '
  'that must be routed through a CCP), never a licence to edit the frozen spec.')
P('ARVES is Universal Cognitive Infrastructure. UCS (Universal Cognitive Standard) is the standard; UCI is its '
  'reference implementation. This manual certifies UCI and any Independent Runtime against UCS.')
table(['Instrument', 'Nature', 'Verdict Space', 'Authority'],
      [['Scenario Conformance Framework', 'Mechanical, executable, property-based', 'PASS / PARTIAL / FAIL per scenario', 'Automated harness + probes'],
       ['Independent Architecture Review', 'Adversarial, human-judged', 'PASS / PARTIAL / FAIL per dimension', 'Review board (arms-length)'],
       ['Certification Levels L1-L4', 'Composite gate over both', 'Level granted / withheld', 'Certification authority'],
       ['Certified Product', 'Product-grade attestation', 'Granted / withheld', 'Certification authority']])

# ---------------- PART 2 ----------------
H('Part 2 - Principles of ARVES Certification')
P('Certification is governed by principles that make the verdict independent, reproducible, and honest.')
B('Property-based, not golden-output. A verdict is a statement about invariants and properties holding across a '
  'trace, NOT a byte-comparison against a recorded expected output. Cognitive systems are non-deterministic at the '
  'inference layer; correctness is defined at the invariant layer.')
B('Independence. The reviewer evaluates a submission as if a rival company built it. Familiarity with the reference '
  'implementation must not soften the verdict.')
B('Reproducibility. Every verdict is replayable from a recorded decision trace (see ORCH-003), not by recomputing '
  'engine inference.')
B('Traceability. Every scenario maps to at least one axis and every node probe maps to a contract clause.')
B('No behaviour without a scenario. CCP-GATE forbids ratifying any behaviour that lacks a conformance scenario.')
B('Certification never edits the spec. Discovered gaps route through CCP, Amendment, IDR, or next major version.')
B('Two planes are respected in every verdict. Truth verdicts are evaluated against the Control Plane / Kernel '
  'boundary; observability data (metrics, logs, tracing, presence) is AP and never treated as truth.')

# ---------------- PART 3 ----------------
H('Part 3 - The Scenario Conformance Framework - Three-Layer Model')
P('Conformance is defined in three separated layers plus a verdict. This separation lets a single axis be exercised '
  'by many scenarios, and a single scenario to be observed by many node probes.')
B('AXIS - a capability dimension the architecture is stressed on (12 defined, see Part 4).')
B('REFERENCE SCENARIO - a concrete, versioned instantiation combining several axes into an executable narrative.')
B('NODE PROBE - an observation point bound to a contract clause; it records what an engine/capability/kernel node '
  'did, so an invariant or property can be evaluated.')
B('VERDICT - PASS / PARTIAL / FAIL, computed from invariant and property checks over the recorded trace.')
table(['Layer', 'Answers', 'Bound To', 'Versioned Against'],
      [['Axis', 'Which dimension of capability is stressed', 'Architecture stress model', 'UCS major version'],
       ['Reference Scenario', 'What concrete workload is run', 'One or more axes', 'Conformance suite version'],
       ['Node Probe', 'What a specific node actually did', 'A contract clause', 'Contract version'],
       ['Verdict', 'Did the invariants/properties hold', 'Registered invariants + properties', 'Both suite and spec version']])

# ---------------- PART 4 ----------------
H('Part 4 - The Twelve Conformance Axes')
P('The twelve axes are the fixed capability dimensions along which any ARVES runtime is stressed. They are frozen '
  'with the specification; new axes require a new UCS major version.')
table(['#', 'Axis', 'What It Stresses', 'Primary Invariant Focus'],
      [['1', 'Information-intensive', 'Large ontology graphs, dense entity/observation load', 'OWN-001, provenance (O-003)'],
       ['2', 'Event-driven', 'Reactive flows over the Event Fabric', 'LAYER-001 cross-cutting via Event Fabric'],
       ['3', 'Human-collaboration', 'Human-in-the-loop decisions, approvals, handoff', 'ORCH-001 (truth only in Kernel)'],
       ['4', 'Multi-step planning', 'Plan/Engine Graph over many steps', 'ORCH-002 (plans, not persistent state)'],
       ['5', 'Long-running', 'Durable state across time, checkpoints', 'ORCH-003 (replay from trace)'],
       ['6', 'Physical-world', 'Actuation, sensing, irreversible effects', 'ORCH-004 (idempotent + content-addressable)'],
       ['7', 'Safety-critical', 'Fail-closed behaviour, no partial truth', 'Amendment-006 failure model'],
       ['8', 'High-volume streaming', 'Sustained throughput, backpressure', 'SHARD-001 partitioning'],
       ['9', 'Multi-agent', 'Many agents coordinating over shared truth', 'ORCH-001..004 under concurrency'],
       ['10', 'Policy-heavy', 'Dense governance/authorization constraints', 'ORCH-001, OWN-001'],
       ['11', 'Autonomous', 'Self-directed goal pursuit, minimal human input', 'ORCH-002, ORCH-003'],
       ['12', 'Recovery/replay', 'Crash, partition, split-brain recovery', 'ORCH-003, IDR-001..005']])
P('Note on invariants: ORCH-001..004, OWN-001, LAYER-001, SHARD-001 are REGISTERED (normative). Ontology principles '
  'O-001..007 are DESIGN PRINCIPLES (definitional, not runtime-provable) and are used to shape scenarios, not to '
  'pass/fail a node. Proposed invariants (G-001, QUERY-001, LCW-001, PERSIST-001, CAP-*, ENG-*) are INFORMATIVE '
  'only and marked proposed (pending CCP) wherever they appear.')

# ---------------- PART 5 ----------------
H('Part 5 - Reference Scenarios')
P('A reference scenario is a concrete, versioned workload combining several axes. Each scenario declares its axis '
  'coverage, its entry conditions, the node probes it activates, and the invariants/properties its verdict evaluates. '
  'Scenarios are illustrative reference workloads, not an exhaustive catalog; the suite grows through CCP-gated '
  'additions.')
table(['Scenario', 'Axes Combined', 'Verdict Properties (examples)'],
      [['Ingest-and-Derive', '1 Information-intensive, 2 Event-driven',
        'Every observation carries provenance; only Kernel commits truth (ORCH-001); derivation != inheritance'],
       ['Plan-and-Act', '4 Multi-step planning, 6 Physical-world',
        'Plans are not persistent state (ORCH-002); every capability invocation idempotent + content-addressable (ORCH-004)'],
       ['Human-Gated Approval', '3 Human-collaboration, 10 Policy-heavy',
        'No effect committed without gate; single owner per state (OWN-001)'],
       ['Long-Run Saga', '5 Long-running, 7 Safety-critical',
        'Partial execution rolled back by NOT committing; committed effects compensated saga-style (Amendment-006)'],
       ['Stream-Under-Load', '8 High-volume streaming, 2 Event-driven',
        'Partition key immutable (SHARD-001); bounded-staleness reads remain within bound'],
       ['Swarm-Coordinate', '9 Multi-agent, 11 Autonomous',
        'Concurrent agents never produce conflicting committed truth; execution replayable (ORCH-003)'],
       ['Crash-and-Replay', '12 Recovery/replay, 5 Long-running',
        'Replay reconstructs from recorded decision trace, NOT recomputation (ORCH-003); WAL = decision trace (IDR-001)']])

# ---------------- PART 6 ----------------
H('Part 6 - Node Probes and Verdict Computation')
P('A node probe is an observation point bound to a contract clause. It records structured evidence about a node - '
  'an engine invocation, a capability binding, a Kernel commit, a Query read - so a property can be evaluated after '
  'the fact. Probes are passive: they observe, they do not steer.')
P('Verdict computation proceeds in three stages:')
B('Collect. Run the scenario; probes append evidence to the recorded decision trace (the WAL / Raft log under IDR-001).')
B('Evaluate. For each declared invariant and property, evaluate a predicate over the collected trace. Predicates are '
  'property-based (e.g. "for all committed effects, exactly one owner") not golden-output comparisons.')
B('Aggregate. Combine per-property results into a scenario verdict.')
table(['Verdict', 'Meaning', 'Gate Effect'],
      [['PASS', 'All declared invariants and properties held across the trace', 'Counts toward level certification'],
       ['PARTIAL', 'Core invariants held; one or more non-blocking properties failed or were unproven', 'Level withheld until resolved; documented as known limitation'],
       ['FAIL', 'A registered invariant was violated', 'Blocks certification; defect in implementation']])
P('A registered-invariant violation is always FAIL. A proposed-invariant expectation (G-001/QUERY-001/LCW-001/'
  'PERSIST-001/CAP-*/ENG-*, all proposed pending CCP) can at most produce PARTIAL, never FAIL, because it is not '
  'yet normative.')

# ---------------- PART 7 ----------------
H('Part 7 - The Conformance Artifact')
P('Every certification run emits a single immutable conformance artifact. It is the evidence package a third party '
  'can independently replay. Because ORCH-003 requires replay from the recorded decision trace and not recomputation, '
  'the artifact is self-sufficient: given the artifact, a verifier reconstructs the verdict without re-running engines.')
table(['Field', 'Content', 'Source'],
      [['Suite version', 'Conformance suite version identifier', 'Suite registry'],
       ['Spec version', 'UCS version under test (frozen 2026-07-01 for v1)', 'Specification Freeze Record'],
       ['Runtime identity', 'Implementation name + build (UCI or Independent Runtime A/B)', 'Submission'],
       ['Axis coverage', 'Which of the 12 axes were exercised', 'Scenario declarations'],
       ['Scenario results', 'Per-scenario PASS/PARTIAL/FAIL', 'Verdict computation'],
       ['Decision trace', 'Recorded WAL / Raft log enabling replay (ORCH-003, IDR-001/005)', 'Runtime under test'],
       ['Invariant matrix', 'Per registered invariant: held / violated', 'Property evaluation'],
       ['Proposed-invariant notes', 'Informative results for proposed invariants, flagged pending CCP', 'Property evaluation'],
       ['Level attestation', 'Which of L1-L4 / Certified Product is supported', 'Aggregation']])

# ---------------- PART 8 ----------------
H('Part 8 - Certification Levels L1-L4 and Certified Product')
P('Certification levels are cumulative. A runtime must hold every lower level before a higher level is granted. '
  'Levels map to the frozen milestones I1..I6 and to the layer/plane responsibilities.')
table(['Level', 'Name', 'What It Certifies', 'Milestone Alignment'],
      [['L1', 'Core Runtime', 'Single-node truth: Kernel owns TRUTH and commits; Engine pure/stateless produces only proposed effects; Query READ-ONLY; OWN-001 holds', 'I1 Distributed Runtime (single-node baseline)'],
       ['L2', 'Cognitive Control', 'Control Plane owns Plan/Engine Graph; ORCH-001/002 hold (no truth, no persistent state in CP); LCW owns working memory not truth', 'I1 -> I2'],
       ['L3', 'Distributed', 'Cluster Kernel: per-shard Raft (IDR-001), replicate committed OUTCOMES not invocations, WAL = decision trace, sagas for cross-shard; SHARD-001 immutable partition key; consistency tiers honored', 'I2 Cluster Kernel, I3 Distributed Query, I4 Capability Scheduling'],
       ['L4', 'Multi-Agent', 'Multi-agent runtime: many agents coordinate over shared committed truth without conflict; ORCH-003/004 under concurrency; recovery/replay axis passes', 'I5 Multi-Agent Runtime'],
       ['Certified Product', 'Product-grade', 'Enterprise-readiness met; real product built on ARVES without modifying the standard; all lower levels held', 'I6 Reference Products']])
P('Kernel-never-Control-Plane rule: at every level, the Kernel commits truth but never becomes the Control Plane. '
  'A submission where the Kernel orchestrates, plans, or executes FAILS regardless of other results.')

# ---------------- PART 9 ----------------
H('Part 9 - The Independent Architecture Review Process')
P('The Independent Architecture Review is the human, adversarial counterpart to the mechanical conformance suite. '
  'The reviewer evaluates the submission as if a competing company had built and submitted it - no benefit of the '
  'doubt, no reliance on shared authorship with the reference implementation.')
P('Each dimension receives an independent verdict of PASS, PARTIAL, or FAIL.')
table(['Review Dimension', 'Question Asked', 'FAIL Trigger'],
      [['Layering', 'Are dependencies downward-only with no lateral coupling (LAYER-001)?', 'Any lateral or upward dependency'],
       ['Ownership', 'Does every state have exactly one owner (OWN-001)?', 'Any shared or ownerless state'],
       ['Plane separation', 'Is Control Plane strictly decide, Data Plane strictly carry?', 'Kernel acting as Control Plane; CP holding truth (ORCH-001)'],
       ['Truth discipline', 'Does only the Kernel commit truth; are engine writes proposed effects only?', 'Engine or Query writing committed truth'],
       ['Orchestration', 'Plans not persistent state (ORCH-002); replay from trace (ORCH-003); idempotent + content-addressable (ORCH-004)?', 'Replay by recomputation; non-idempotent invocation'],
       ['Distribution', 'Does it follow IDR-001..005 (per-shard Raft, replicate outcomes, sagas, no cross-shard atomic commit in v1)?', 'Cross-shard atomic commit claimed in v1'],
       ['Consistency', 'Are linearizable / bounded-staleness / eventual tiers correctly implemented?', 'Stale read presented as linearizable'],
       ['Failure handling', 'Partial rollback by NOT committing; saga compensation; cooperative idempotent cancellation (Amendments 005/006)?', 'Partial truth left committed after failure'],
       ['Ontology fidelity', 'Do types honor O-001..007 design principles (identity, provenance, versioned+registered)?', 'Unversioned/unregistered types; derivation used as inheritance'],
       ['Conformance integrity', 'Do scenarios and probes trace to axes and contract clauses; CCP-GATE respected?', 'Behaviour claimed without a conformance scenario']])
P('Review procedure:')
B('Intake - receive submission, runtime identity, and conformance artifact.')
B('Blind pass - reviewer reconstructs the architecture from artifacts alone, without designer narration.')
B('Adversarial probing - reviewer attempts to construct a scenario that violates a registered invariant.')
B('Dimension scoring - assign PASS/PARTIAL/FAIL per dimension with cited evidence.')
B('Disposition - overall PASS requires all dimensions PASS; any FAIL blocks; PARTIALs are documented as conditions.')
B('Routing - genuine spec gaps discovered during review are routed to CCP / Amendment / IDR, never patched into the spec.')

# ---------------- PART 10 ----------------
H('Part 10 - Conformance Checklist')
P('Use this checklist to gate a runtime before submitting for level certification. Every item must be demonstrable '
  'from the conformance artifact.')
B('All 12 axes have at least one reference scenario exercised for the target level.')
B('Every scenario verdict is PASS or an explicitly documented PARTIAL; no unexplained FAIL.')
B('OWN-001: every state in the trace resolves to exactly one owner.')
B('ORCH-001: no truth committed outside the Kernel; Control Plane holds no truth.')
B('ORCH-002: no persistent state produced by the Control Plane; only plans.')
B('ORCH-003: verdict replayable from recorded decision trace, not recomputation.')
B('ORCH-004: every engine/capability invocation idempotent and content-addressable.')
B('LAYER-001: dependency graph is downward-only; cross-cutting via Control Plane / Event Fabric only.')
B('SHARD-001: partition by tenant/workspace; partition key immutable across the trace.')
B('Query nodes wrote nothing (READ-ONLY).')
B('Engine nodes are pure/stateless; all writes appear as proposed effects the Kernel later commits.')
B('Proposed invariants (G-001/QUERY-001/LCW-001/PERSIST-001/CAP-*/ENG-*) reported as informative only, flagged pending CCP.')
B('Conformance artifact is complete, immutable, and independently replayable.')
B('Suite version and spec version recorded and compatible (see Part 14).')

# ---------------- PART 11 ----------------
H('Part 11 - Performance and Benchmark Checklist')
P('Performance is measured as properties over the same recorded traces, never as truth. Benchmarks stress the '
  'high-volume streaming, long-running, and distributed axes.')
table(['Metric', 'Axis Stressed', 'Property Checked'],
      [['Commit latency (leader)', 'Safety-critical, Distributed', 'Linearizable commit through shard leader within target'],
       ['Bounded-staleness read', 'High-volume streaming', 'Follower reads within declared staleness bound'],
       ['Eventual read convergence', 'Event-driven', 'Replica converges within target window'],
       ['Throughput under backpressure', 'High-volume streaming', 'No dropped commits; backpressure applied, not silent loss'],
       ['Replay time', 'Recovery/replay', 'Trace replays without re-running engines (ORCH-003)'],
       ['Recovery after crash', 'Recovery/replay', 'Leader re-election + WAL replay restores truth (IDR-001/004/005)'],
       ['Saga compensation time', 'Long-running, Safety-critical', 'Committed effects compensated within target (Amendment-006)']])
P('Benchmark checklist:')
B('Observability data (metrics/logs/tracing/presence) is treated as AP, never as committed truth.')
B('Every benchmark run emits its own decision trace so results are replayable.')
B('Load is partitioned by immutable partition key (SHARD-001); cross-shard work uses sagas, not atomic commit.')
B('Latency is reported per consistency tier (linearizable / bounded-staleness / eventual), never conflated.')
B('No benchmark relies on golden-output comparison; all pass/fail is property-based.')

# ---------------- PART 12 ----------------
H('Part 12 - Enterprise-Readiness Checklist')
P('Enterprise readiness is a precondition for the Certified Product attestation. It certifies that a runtime is '
  'operable, governable, and safe in production, on top of holding L1-L4.')
B('Multi-tenant isolation enforced by immutable partition key (SHARD-001); no cross-tenant truth leakage.')
B('Read consistency tiers exposed and documented (linearizable through leader, bounded-staleness follower, eventual replica).')
B('Failure model implemented: partial execution rolled back by NOT committing; committed effects compensated saga-style (Amendment-006).')
B('Cancellation is cooperative and idempotent; no partial truth; priority/preemption via checkpoint -> replay (Amendment-005).')
B('Cluster membership managed via Raft joint consensus (IDR-003); per-shard leader election (IDR-004).')
B('Append-only WAL + snapshots for durability and replay (IDR-005); WAL doubles as decision trace.')
B('Governance/policy constraints enforced at the Control Plane decision boundary, never bypassed by the Data Plane.')
B('Every type is versioned and registered (O-006); ontology defines meaning, not storage (O-007).')
B('Observability stack (AP) deployed and separated from truth (CP).')
B('Upgrade path preserves conformance: new build re-passes the compatible conformance suite before promotion.')
B('SDK and API surface do not permit clients to write committed truth directly (only through Kernel commit path).')

# ---------------- PART 13 ----------------
H('Part 13 - Independent Runtime A/B Goals')
P('A core certification goal is that ARVES is not a single-vendor artifact. Two independent runtimes - Independent '
  'Runtime A and Independent Runtime B - must each pass certification against the same frozen UCS, proving the '
  'standard is implementable by more than its authors.')
table(['Runtime', 'Goal', 'Success Criterion'],
      [['UCI (reference)', 'Prove the specification is implementable', 'Holds L1-L4 + Certified Product'],
       ['Independent Runtime A', 'Prove standard is vendor-independent', 'Passes certification independently against same spec version'],
       ['Independent Runtime B', 'Prove interoperability of the standard', 'Passes certification; produces artifacts replayable by A and UCI verifiers']])
P('A/B parity properties:')
B('Both runtimes accept the same conformance suite version against the same UCS version.')
B('Both emit conformance artifacts that are cross-verifiable (a verifier built for one can replay the other via the recorded trace).')
B('Neither runtime requires modification of the frozen standard to pass; any gap routes through CCP.')
B('Distributed behaviour of both grounds in IDR-001..005 (per-shard Raft, replicate outcomes, sagas), even if engine internals differ.')

# ---------------- PART 14 ----------------
H('Part 14 - Versioning Conformance Suites Against Spec Versions')
P('A verdict is only meaningful when the suite version and the spec version are compatible. The Specification Era is '
  'frozen as of 2026-07-01; the conformance suite evolves under CCP-GATE without ever mutating the frozen spec.')
table(['Artifact', 'Maturity Lifecycle', 'Change Instrument'],
      [['UCS specification', 'Draft -> Candidate -> Ratified -> Frozen -> Deprecated -> Superseded', 'CCP / Amendment / IDR / next major version'],
       ['Conformance suite', 'Draft -> Candidate -> Ratified -> Frozen', 'CCP (CCP-GATE: no behaviour ratified without a scenario)'],
       ['Contract (probe binding)', 'Draft -> Candidate -> Ratified -> Frozen', 'CCP / Amendment'],
       ['Certification artifact', 'Immutable once emitted', 'Re-run produces a new artifact']])
P('Versioning rules:')
B('A conformance suite declares the exact UCS version it certifies against; certifying across incompatible versions is invalid.')
B('Adding a scenario or probe is a suite minor version; it never changes the spec.')
B('Adding or removing an axis requires a new UCS major version (the 12 axes are frozen for v1).')
B('Promoting a proposed invariant (G-001/QUERY-001/LCW-001/PERSIST-001/CAP-*/ENG-*) to normative requires a ratified '
  'CCP; only then may a related expectation escalate from PARTIAL to FAIL.')
B('Reference Lifecycle context: Conformance (stage) precedes Reference Runtime, Certification, Reference Product, and '
  'Reference Ecosystem; certification cannot outrun the lifecycle stage it depends on.')

# ---------------- PART 15 ----------------
H('Part 15 - Certified Ecosystem and Product Goals')
P('The terminal goal of certification is a living ecosystem: multiple certified runtimes, real products, and a '
  'marketplace, all built on ARVES without modifying the standard. These are the frozen ecosystem goals.')
B('Production distributed runtime certified at L3+.')
B('Complete conformance suite covering all 12 axes.')
B('Independent Runtime A and Independent Runtime B both pass certification.')
B('Third-party certification available (arms-length review board, not the reference authors).')
B('Enterprise runtime meeting the enterprise-readiness checklist and Certified Product bar.')
B('SDKs that cannot bypass the Kernel commit path.')
B('Marketplace of certified capabilities and products.')
B('Cloud offering exposing documented consistency tiers.')
B('Real products built on ARVES without modifying the standard (I6 Reference Products).')
P('These goals align to the frozen milestones I1 Distributed Runtime, I2 Cluster Kernel, I3 Distributed Query, '
  'I4 Capability Scheduling, I5 Multi-Agent Runtime, I6 Reference Products - and to no others.')

# ---------------- PART 16 ----------------
H('Part 16 - Certification Decision Matrix and Sign-Off')
P('The final certification decision composes the mechanical conformance verdict with the Independent Architecture '
  'Review verdict. Both must be satisfied for a level to be granted.')
table(['Conformance Result', 'Architecture Review', 'Decision'],
      [['All target-level scenarios PASS', 'All dimensions PASS', 'Level GRANTED'],
       ['Some PARTIAL (documented, non-blocking)', 'All dimensions PASS', 'Level GRANTED with documented conditions'],
       ['Any registered-invariant FAIL', 'Any', 'Level WITHHELD (implementation defect)'],
       ['All PASS', 'Any dimension FAIL', 'Level WITHHELD (architecture defect)'],
       ['Any', 'Spec gap discovered', 'Route to CCP; do NOT edit spec; re-certify after resolution']])
P('Sign-off record fields:')
B('Runtime identity and build; UCS version; conformance suite version.')
B('Level attested (L1 / L2 / L3 / L4 / Certified Product).')
B('Conformance artifact reference (immutable, replayable).')
B('Architecture review dimension scores with cited evidence.')
B('Documented PARTIAL conditions and their remediation owners.')
B('Any CCP routed as a result of the review.')

# ---------------- FINAL DEFINITION ----------------
I('Final Definition  Certification is the independent, replayable proof that an implementation earns the ARVES name - '
  'PASS/PARTIAL/FAIL verdicts drawn from invariants and properties, never golden output - and it proves the frozen '
  'standard without ever changing it.')

d.save(FN)
print(FN)
